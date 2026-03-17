import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Define paths for images (using the IDs from the previous tool outputs)
    brain_dir = r"C:\Users\Administrator\.gemini\antigravity\brain\b34e3501-2ed3-4089-bfcd-c2274a685cc8"
    img_goals = os.path.join(brain_dir, "goal_hierarchy_diagram_1773341332729.png")
    img_caps = os.path.join(brain_dir, "capability_grouping_diagram_1773341348464.png")
    img_acq = os.path.join(brain_dir, "acquaintance_diagram_1773341363082.png")
    img_int = os.path.join(brain_dir, "interaction_diagram_1773341482787.png")

    # Title Page
    title = doc.add_heading('DCIT 403 — Semester Project Report\nUniversity of Ghana', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Crop Disease Intelligent Agent', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # TOC
    doc.add_heading('TABLE OF CONTENTS', 1)
    toc_topics = [
        "1. Introduction", "2. System Analysis", "3. Architectural Design", 
        "4. Interaction Design", "5. Detailed Design", "6. Implementation", 
        "7. Conclusion", "8. References"
    ]
    for topic in toc_topics:
        doc.add_paragraph(topic)
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. INTRODUCTION', 1)
    doc.add_heading('1.1 Problem Statement', 2)
    doc.add_paragraph(
        "Crop diseases are responsible for an estimated 10–16% annual loss in global agricultural production, "
        "costing billions of dollars and threatening food security. Smallholder farmers—who produce over 70% "
        "of the world's food—are disproportionately affected because they lack timely access to expert plant "
        "pathologists. By the time a farmer visually identifies a disease, significant yield damage may have already occurred.\n\n"
        "The Crop Disease Intelligent Agent addresses this by providing instant, AI-powered diagnosis of crop "
        "diseases from leaf images."
    )

    doc.add_heading('1.2 Why an Agent-Based Approach?', 2)
    doc.add_paragraph(
        "An agent-based approach is ideal for this problem because the system must demonstrate autonomy, "
        "reactivity, proactiveness, and goal-directed behaviour—properties that distinguish agents from traditional software applications."
    )
    
    # Add a table for agent properties
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Agent Property'
    hdr_cells[1].text = 'How It Applies'
    
    properties = [
        ("Autonomy", "The agent independently processes images and produces diagnoses without human intervention"),
        ("Reactivity", "It reacts to new image inputs (percepts) from the environment in real time"),
        ("Proactiveness", "It pursues the goal of accurate diagnosis and can escalate uncertain cases"),
        ("Goal-directed", "It has clear goals: maximise diagnostic accuracy, provide actionable treatment plans, and flag uncertain results")
    ]
    for prop, app in properties:
        row_cells = table.add_row().cells
        row_cells[0].text = prop
        row_cells[1].text = app

    # 2. System Analysis
    doc.add_heading('2. SYSTEM ANALYSIS', 1)
    doc.add_heading('2.1 Goal Specification', 2)
    doc.add_paragraph("The agent is driven by top-level goals and sub-goals structured in a hierarchy.")
    
    if os.path.exists(img_goals):
        doc.add_picture(img_goals, width=Inches(6))
        doc.add_paragraph("Figure 1: Goal Hierarchy Diagram").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('2.2 Functionalities', 2)
    doc.add_paragraph("The system provides Image Intake, Disease Identification, Confidence Reporting, and Field Monitoring.")

    doc.add_heading('2.3 Usage Scenarios', 2)
    doc.add_paragraph("Scenario 1: Farmer uploads a diseased leaf for diagnosis.")
    doc.add_paragraph("Scenario 2: Healthy crop confirmation.")
    doc.add_paragraph("Scenario 3: Low-confidence results escalated to an expert.")

    # 3. Architectural Design
    doc.add_heading('3. ARCHITECTURAL DESIGN', 1)
    doc.add_heading('3.1 Agent Type Selection', 2)
    doc.add_paragraph("The system uses a single BDI agent with specialized capabilities.")

    doc.add_heading('3.2 Capability Grouping', 2)
    if os.path.exists(img_caps):
        doc.add_picture(img_caps, width=Inches(6))
        doc.add_paragraph("Figure 2: Functionality Grouping Diagram").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('3.3 Acquaintance Diagram', 2)
    if os.path.exists(img_acq):
        doc.add_picture(img_acq, width=Inches(6))
        doc.add_paragraph("Figure 3: Acquaintance Diagram").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4. Interaction Design
    doc.add_heading('4. INTERACTION DESIGN', 1)
    doc.add_heading('4.1 Interaction Diagrams', 2)
    if os.path.exists(img_int):
        doc.add_picture(img_int, width=Inches(6))
        doc.add_paragraph("Figure 4: Sequence Interaction Diagram").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5. Detailed Design
    doc.add_heading('5. DETAILED DESIGN', 1)
    doc.add_heading('5.1 Capabilities', 2)
    doc.add_paragraph("Perception, Classification, Treatment Recommendation, and Monitoring.")
    
    doc.add_heading('5.2 Plans', 2)
    doc.add_paragraph("DiagnoseSingleImage, HandleLowConfidence, BatchMonitoring, and HandleUnknownDisease.")

    # 6. Implementation
    doc.add_heading('6. IMPLEMENTATION', 1)
    doc.add_paragraph(
        "The agent is implemented in Python using PyTorch for the CNN model and Gradio for the UI. "
        "A full BDI loop (Perceive-Decide-Act) is used for decision making."
    )

    # 7. Conclusion
    doc.add_heading('7. CONCLUSION', 1)
    doc.add_paragraph(
        "The Crop Disease Intelligent Agent successfully demonstrates the application of the Prometheus "
        "methodology to a real-world problem, providing farmers with autonomous and reliable diagnostic tools."
    )

    # Save the document
    desktop_path = os.path.join(os.environ['USERPROFILE'], 'Desktop')
    output_path = os.path.join(desktop_path, "Crop_Disease_Agent_Project_Report.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == "__main__":
    create_report()
